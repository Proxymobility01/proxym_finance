from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

from django.conf import settings

def _compose_title_from_request(request) -> str:
    """
    Titre demandé :
    'RECAPITULATIF DU <date> LEASES PAYE AVANT 14H'
    - si ?date_concernee=YYYY-MM-DD => utilise cette date
    - si range => 'RECAPITULATIF DU <start> AU <end> LEASES PAYE AVANT 14H'
    - sinon => 'RECAPITULATIF LEASES PAYE AVANT 14H'
    """
    d  = (request.GET.get("date_concernee") or "").strip()
    da = (request.GET.get("date_concernee_after") or "").strip()
    db = (request.GET.get("date_concernee_before") or "").strip()

    if d:
        return f"RECAPITULATIF DU {d} LEASES PAYE AVANT 14H"
    elif da and db:
        return f"RECAPITULATIF DU {da} AU {db} LEASES PAYE AVANT 14H"
    elif da and not db:
        return f"RECAPITULATIF A PARTIR DU {da} LEASES PAYE AVANT 14H"
    elif db and not da:
        return f"RECAPITULATIF JUSQU'AU {db} LEASES PAYE AVANT 14H"
    return "RECAPITULATIF LEASES PAYE AVANT 14H"

def _load_template_document() -> Document:
    """
    Charge le modèle Word; si absent, retourne un Document vierge.
    """
    tpl = Path(settings.BASE_DIR) / "templates"  / "rapport_leases.docx"
    if tpl.exists():
        return Document(str(tpl))
    return Document()

def _replace_placeholders_everywhere(doc: Document, mapping: dict[str, str]) -> None:
    """
    Remplace {{PLACEHOLDER}} dans :
    - paragraphes,
    - en-têtes / pieds,
    - cellules de tableaux.
    """
    def _replace_in_paragraphs(paragraphs):
        for p in paragraphs:
            for key, val in mapping.items():
                if key in p.text:
                    # reconstruire le run pour préserver le style proprement
                    inline = p.runs
                    text = p.text.replace(key, val)
                    # clear
                    for i in range(len(inline)-1, -1, -1):
                        p.runs[i].clear()
                    p.text = text

    # Corps
    _replace_in_paragraphs(doc.paragraphs)

    # Tableaux
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _replace_in_paragraphs(cell.paragraphs)

    # En-têtes/pieds
    for section in doc.sections:
        _replace_in_paragraphs(section.header.paragraphs)
        _replace_in_paragraphs(section.footer.paragraphs)
